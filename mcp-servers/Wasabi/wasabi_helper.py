import boto3
from io import BytesIO
import csv
import os
import tempfile
import subprocess
from pypdf import PdfReader
from docx import Document
import pandas as pd
import xlrd
import openpyxl
import config

def get_s3_object_body(bucket_name, object_key):
    """
    Helper function to get the raw binary content of an S3 object.
    """
    try:
        s3_client = boto3.client(
            's3',
            endpoint_url = config.WASABI_ENDPOINT_URL,
            aws_access_key_id = config.WASABI_ACCESS_KEY,
            aws_secret_access_key = config.WASABI_SECRET_KEY
        )            

        response = s3_client.get_object(Bucket=bucket_name, Key=object_key)
        return response['Body'].read()
    except s3_client.exceptions.NoSuchKey:
        print(f"Error: Object '{object_key}' not found in bucket '{bucket_name}'.")
        return None
    except Exception as e:
        print(f"An S3 error occurred: {e}")
        return None

def read_pdf_from_s3(bucket_name, object_key, **kwargs):
    """
    Reads a PDF file from S3 and returns its text content.
    Requires 'pypdf' or 'PyPDF2' to be installed.
    """
    if PdfReader is None:
        print("Error: 'pypdf' library is not installed. Cannot read PDF files.")
        return None

    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    try:
        pdf_file = BytesIO(body)
        reader = PdfReader(pdf_file)
        text_content = ""
        for page in reader.pages:
            text_content += page.extract_text() + "\n"
        return text_content
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def read_docx_from_s3(bucket_name, object_key, **kwargs):
    """
    Reads a DOCX file from S3 and returns its text content.
    Requires 'python-docx' to be installed.
    """
    if Document is None:
        print("Error: 'python-docx' library is not installed. Cannot read DOCX files.")
        return None

    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    try:
        document_file = BytesIO(body)
        document = Document(document_file)
        full_text = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        cell_text.append(paragraph.text)
                    row_text.append(" ".join(cell_text))
                full_text.append("\t".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def read_doc_from_s3(bucket_name, object_key, **kwargs):
    """
    Reads a DOC file from S3 by attempting to use 'antiword' to convert it
    to text. Requires 'antiword' to be installed on the system where this
    script is executed.
    """
    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    temp_file_path = None
    try:
        # Create a temporary file to save the .doc content
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp_file:
            tmp_file.write(body)
            temp_file_path = tmp_file.name

        try:
            # Use antiword to convert .doc to text
            result = subprocess.run(
                ['antiword', temp_file_path],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except FileNotFoundError:
            print("Error: 'antiword' not found. Please install antiword on your system (e.g., sudo apt-get install antiword on Linux).")
            return None
        except subprocess.CalledProcessError as e:
            print(f"Error converting .doc with antiword: {e.stderr}")
            return None
    except Exception as e:
        print(f"An unexpected error occurred reading DOC: {e}")
        return None
    finally:
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path) # Clean up the temporary file

def read_txt_from_s3(bucket_name, object_key, encoding='utf-8', **kwargs):
    """
    Reads a TXT file from S3 and returns its content as a string.
    """
    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    try:
        return body.decode(encoding)
    except UnicodeDecodeError:
        print(f"Warning: Could not decode TXT with {encoding}. Trying 'latin-1'.")
        return body.decode('latin-1') # Fallback for common encoding issues
    except Exception as e:
        print(f"Error reading TXT: {e}")
        return None

def read_csv_from_s3(bucket_name, object_key, encoding='utf-8', delimiter=',', as_list=False, **kwargs):
    """
    Reads a CSV file from S3.
    Args:
        as_list (bool): If True, returns a list of lists. Otherwise, returns raw string.
    """
    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    try:
        csv_string = body.decode(encoding)
        if as_list:
            reader = csv.reader(csv_string.splitlines(), delimiter=delimiter)
            return list(reader)
        else:
            return csv_string
    except UnicodeDecodeError:
        print(f"Warning: Could not decode CSV with {encoding}. Trying 'latin-1'.")
        csv_string = body.decode('latin-1')
        if as_list:
            reader = csv.reader(csv_string.splitlines(), delimiter=delimiter)
            return list(reader)
        else:
            return csv_string
    except Exception as e:
        print(f"Error reading CSV: {e}")
        return None

def read_xls_from_s3(bucket_name, object_key, sheet_name=0, **kwargs):
    """
    Reads an XLS file from S3 and returns the content of a specified sheet as a string
    (by converting DataFrame to string). Requires 'pandas' and 'xlrd' to be installed.
    """
    if pd is None or xlrd is None:
        print("Error: 'pandas' and/or 'xlrd' libraries are not installed. Cannot read XLS files.")
        return None

    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    try:
        xls_file_stream = BytesIO(body)
        df = pd.read_excel(xls_file_stream, sheet_name=sheet_name, engine='xlrd')
        return df.to_string(index=False) # Convert DataFrame to string, no index
    except Exception as e:
        print(f"Error reading XLS: {e}")
        return None

def read_xlsx_from_s3(bucket_name, object_key, sheet_name=0, **kwargs):
    """
    Reads an XLSX file from S3 and returns the content of a specified sheet as a string
    (by converting DataFrame to string). Requires 'pandas' and 'openpyxl' to be installed.
    """
    if pd is None or openpyxl is None:
        print("Error: 'pandas' and/or 'openpyxl' libraries are not installed. Cannot read XLSX files.")
        return None

    body = get_s3_object_body(bucket_name, object_key, **kwargs)
    if body is None:
        return None

    try:
        xlsx_file_stream = BytesIO(body)
        df = pd.read_excel(xlsx_file_stream, sheet_name=sheet_name, engine='openpyxl')
        return df.to_string(index=False) # Convert DataFrame to string, no index
    except Exception as e:
        print(f"Error reading XLSX: {e}")
        return None